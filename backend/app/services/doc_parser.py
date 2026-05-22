"""
文档解析服务：提取 DOCX/PDF 文本，结构化为条款列表。
支持 DOCX 内嵌图片自动 OCR（混合文档场景）。
"""

import os
import re
import logging
import zipfile
import tempfile
from pathlib import Path
from typing import Optional, List

logger = logging.getLogger(__name__)


def _extract_images_from_docx(file_path: str) -> List[str]:
    """
    从 DOCX 中提取所有内嵌图片，保存为临时文件。
    返回图片路径列表。
    """
    image_paths = []
    try:
        with zipfile.ZipFile(file_path) as z:
            media_files = [f for f in z.namelist() if f.startswith('word/media/')]
            for mf in media_files:
                data = z.read(mf)
                ext = Path(mf).suffix or '.png'
                tmp = tempfile.NamedTemporaryFile(suffix=ext, delete=False)
                tmp.write(data)
                tmp.close()
                image_paths.append(tmp.name)
                logger.info(f"Extracted image: {mf} -> {tmp.name} ({len(data)}B)")
    except Exception as e:
        logger.warning(f"Failed to extract images from DOCX: {e}")
    return image_paths


async def _ocr_images(image_paths: List[str]) -> str:
    """
    对图片列表进行 OCR，返回合并后的文本。
    """
    from .ocr_service import ocr_image
    
    texts = []
    for img_path in image_paths:
        try:
            text = await ocr_image(img_path)
            if text and text.strip():
                texts.append(text.strip())
                logger.info(f"OCR image {img_path}: {len(text)} chars")
        except Exception as e:
            logger.warning(f"OCR failed for {img_path}: {e}")
        finally:
            # 清理临时文件
            try:
                os.unlink(img_path)
            except:
                pass
    
    return "\n\n".join(texts)


def extract_text_from_docx(file_path: str) -> str:
    """从 DOCX 文件提取纯文本（不含图片内容）。"""
    from docx import Document

    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                paragraphs.append(" | ".join(row_texts))

    return "\n\n".join(paragraphs)


def extract_text_from_pdf(file_path: str) -> str:
    """从 PDF 文件提取纯文本，保留页码标记。"""
    from PyPDF2 import PdfReader

    reader = PdfReader(file_path)
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text()
        if text:
            pages.append(f"--- PAGE {i} ---\n{text.strip()}")
    return "\n\n".join(pages)


def extract_text(file_path: str) -> str:
    """根据文件类型提取文本。"""
    ext = Path(file_path).suffix.lower()
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".doc":
        logger.warning(".doc format may not be fully supported, attempting .docx parsing")
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


def _page_at_offset(offset: int, page_map: list[tuple[int, int, int]]) -> int:
    """
    根据 page_map 找到给定字符偏移量所在的页码。
    page_map: [(page_num, start_offset, end_offset), ...]
    """
    for page_num, start, end in page_map:
        if start <= offset < end:
            return page_num
    # 超出范围则返回最后一页
    return page_map[-1][0] if page_map else 0


def _build_page_map(text: str) -> list[tuple[int, int, int]]:
    """
    扫描文本中的 --- PAGE N --- 标记，构建 (页码, 起始偏移, 结束偏移) 映射。
    返回 [(page_num, start_offset, end_offset), ...]
    """
    page_marker_re = re.compile(r'---\s*PAGE\s+(\d+)\s*---')
    page_positions = []
    for m in page_marker_re.finditer(text):
        page_num = int(m.group(1))
        page_positions.append((page_num, m.start()))

    if not page_positions:
        return []

    page_map = []
    for i, (page_num, start) in enumerate(page_positions):
        end = page_positions[i + 1][1] if i + 1 < len(page_positions) else len(text)
        page_map.append((page_num, start, end))
    return page_map


def split_into_clauses(text: str) -> list[dict]:
    """
    将合同文本拆分为条款列表。
    返回 [{"title": "...", "content": "...", "page_start": N, "page_end": N}]
    如果文本中包含 --- PAGE N --- 标记（由 extract_text_from_pdf 插入），
    会自动标注每个条款的页码范围。
    """
    # Build page map from markers
    page_map = _build_page_map(text)

    # Strip page markers for clause splitting (but keep text positions accurate)
    # We'll work on original text so offsets are correct
    clauses = []

    # 常见的中文合同条款编号模式
    patterns = [
        r'(?:^|\n)(第[一二三四五六七八九十百千万\d]+[条章节款项])\s*',
        r'(?:^|\n)((?:[一二三四五六七八九十百千万]+)、)\s*',
        r'(?:^|\n)(\d+[.、)）])\s*',
        r'(?:^|\n)((?:Article|Section|Clause)\s+\d+[.]?\d*)\s*',
        r'(?:^|\n)([（(][一二三四五六七八九十\d]+[）)])\s*',
    ]

    combined = '|'.join(f'({p})' for p in patterns)
    matches = list(re.finditer(combined, text, re.MULTILINE | re.IGNORECASE))

    # Filter out matches that are inside page markers
    page_marker_re = re.compile(r'---\s*PAGE\s+\d+\s*---')
    page_marker_spans = [m.span() for m in page_marker_re.finditer(text)]
    def _is_page_marker(match):
        return any(ms <= match.start() < me for ms, me in page_marker_spans)
    matches = [m for m in matches if not _is_page_marker(m)]

    if len(matches) >= 2:
        for i, match in enumerate(matches):
            start = match.start()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            clause_text = text[start:end].strip()
            # Remove any page markers within clause text
            clause_text = page_marker_re.sub('', clause_text).strip()
            lines = clause_text.split('\n', 1)
            title = lines[0].strip()[:100]
            c = lines[1].strip() if len(lines) > 1 else ""
            if c or title:
                entry = {"title": title, "content": c}
                if page_map:
                    entry["page_start"] = _page_at_offset(start, page_map)
                    entry["page_end"] = _page_at_offset(end - 1, page_map)
                clauses.append(entry)
    else:
        # Fallback: split by double newlines
        paragraphs = re.split(r'\n\s*\n', text)
        for i, para in enumerate(paragraphs):
            para = para.strip()
            # Skip page markers
            if page_marker_re.match(para):
                continue
            para = page_marker_re.sub('', para).strip()
            if not para:
                continue
            lines = para.split('\n', 1)
            title = lines[0].strip()[:100]
            c = lines[1].strip() if len(lines) > 1 else ""
            entry = {
                "title": title if len(paragraphs) > 1 else f"段落 {i + 1}",
                "content": c if c else title,
            }
            # We don't have reliable offsets in fallback mode, skip page annotation
            clauses.append(entry)

    return clauses


def parse_document(file_path: str) -> dict:
    """
    解析文档，返回结构化结果。
    """
    ext = Path(file_path).suffix.lower()
    raw_text = extract_text(file_path)
    clauses = split_into_clauses(raw_text)

    logger.info(f"Parsed {file_path}: {len(raw_text)} chars, {len(clauses)} clauses")

    return {
        "raw_text": raw_text,
        "clauses": clauses,
        "total_clauses": len(clauses),
        "file_type": ext,
    }


async def smart_parse_document(file_path: str) -> dict:
    """
    智能文档解析：自动选择文本提取或OCR。
    - DOCX：提取文字 + 提取内嵌图片并 OCR，合并结果
    - PDF：文字提取，不足则 OCR
    - 图片：直接 OCR
    """
    from .ocr_service import smart_extract

    ext = Path(file_path).suffix.lower()

    # 图片直接走 OCR
    if ext in (".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".tiff", ".tif"):
        raw_text = await smart_extract(file_path)
        clauses = split_into_clauses(raw_text)
        logger.info(f"OCR parsed image: {len(raw_text)} chars, {len(clauses)} clauses")
        return {
            "raw_text": raw_text,
            "clauses": clauses,
            "total_clauses": len(clauses),
            "file_type": ext,
            "ocr_used": True,
        }

    # PDF - 检查是否需要 OCR
    if ext == ".pdf":
        raw_text = extract_text(file_path)
        ocr_used = False
        if len(raw_text.strip()) < 100:
            # 扫描件，走 OCR
            raw_text = await smart_extract(file_path)
            ocr_used = True
        clauses = split_into_clauses(raw_text)
        return {
            "raw_text": raw_text,
            "clauses": clauses,
            "total_clauses": len(clauses),
            "file_type": ext,
            "ocr_used": ocr_used,
        }

    # DOCX - 提取文字 + 提取内嵌图片 OCR
    raw_text = extract_text(file_path)
    
    # 提取内嵌图片并 OCR
    image_paths = _extract_images_from_docx(file_path)
    ocr_text = ""
    ocr_used = False
    if image_paths:
        logger.info(f"Found {len(image_paths)} embedded images in DOCX, running OCR...")
        ocr_text = await _ocr_images(image_paths)
        if ocr_text:
            ocr_used = True
            # 合并：文字在前，OCR内容在后
            raw_text = raw_text + "\n\n--- 手写补充/扫描区域 OCR 识别内容 ---\n\n" + ocr_text
            logger.info(f"DOCX with OCR: text={len(raw_text)} chars (ocr added {len(ocr_text)} chars)")
    
    clauses = split_into_clauses(raw_text)
    return {
        "raw_text": raw_text,
        "clauses": clauses,
        "total_clauses": len(clauses),
        "file_type": ext,
        "ocr_used": ocr_used,
        "ocr_images_count": len(image_paths),
    }
