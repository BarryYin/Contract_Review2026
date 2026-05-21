"""
DOCX Track Changes 导出服务。
T-RO-05: 在原始 DOCX 中以可视化 Track Changes 样式标注采纳的修改建议。

策略说明：
  python-docx 不支持原生 OOXML Track Changes（<w:del>/<w:ins>），
  因此采用「模拟 Track Changes」方案：
    - 原文片段 → 红色 + 删除线（run.font.strike = True）
    - 建议片段 → 蓝色 + 下划线（run.font.underline = True）
  用户在 Word 中可以直观区分删除/插入内容。
"""

import io
import logging
from copy import deepcopy
from typing import List, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# 颜色常量
# ---------------------------------------------------------------------------

COLOR_DELETED = RGBColor(0xFF, 0x00, 0x00)   # 红色 — 原文（删除）
COLOR_INSERTED = RGBColor(0x00, 0x00, 0xFF)  # 蓝色 — 建议（插入）


# ---------------------------------------------------------------------------
# Core export function
# ---------------------------------------------------------------------------

def export_track_changes(
    original_file_path: str,
    adopted_issues: List[dict],
) -> bytes:
    """在原始 DOCX 文件中标注 Track Changes 样式修改并返回新文件的字节。

    Args:
        original_file_path: 原始 DOCX 文件路径。
        adopted_issues: 采纳的问题列表，每个元素需包含
            ``modification_example`` 字典，含 ``original`` 与 ``suggested`` 键。

    Returns:
        修改后的 DOCX 文件 bytes。
    """
    doc = Document(original_file_path)

    # 过滤出有有效 modification_example 的问题
    valid_issues = []
    for issue in adopted_issues:
        mod = issue.get("modification_example")
        if not mod or not isinstance(mod, dict):
            continue
        original = mod.get("original", "").strip()
        suggested = mod.get("suggested", "").strip()
        if original and suggested:
            valid_issues.append({
                "original": original,
                "suggested": suggested,
                "issue_id": issue.get("id", ""),
            })

    if not valid_issues:
        logger.info("No valid adopted issues with modification_example; returning original doc.")
        return _doc_to_bytes(doc)

    applied_count = 0
    for issue in valid_issues:
        if _apply_track_change_to_doc(doc, issue["original"], issue["suggested"]):
            applied_count += 1

    logger.info(
        f"Track Changes export: {applied_count}/{len(valid_issues)} "
        f"issues applied in document."
    )

    return _doc_to_bytes(doc)


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _doc_to_bytes(doc: Document) -> bytes:
    """将 Document 序列化为 bytes。"""
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _apply_track_change_to_doc(
    doc: Document,
    original_text: str,
    suggested_text: str,
) -> bool:
    """在整个文档中搜索 original_text，首次命中处执行 Track Change 替换。

    遍历文档所有段落（含表格中的段落），找到包含目标文本的段落后，
    将其 runs 拆分为：before + deleted-run(红/删) + inserted-run(蓝/下划) + after。

    Returns:
        True 表示成功应用至少一处替换。
    """
    # 收集所有段落（含表格单元格内的段落）
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    for para in all_paragraphs:
        if _try_replace_in_paragraph(para, original_text, suggested_text):
            return True

    logger.debug(f"Track Change target not found in document: '{original_text[:50]}...'")
    return False


def _try_replace_in_paragraph(
    para,
    original_text: str,
    suggested_text: str,
) -> bool:
    """在单个段落中尝试查找并替换一次 original_text。

    由于一个段落的文本可能跨多个 runs，我们需要：
    1. 拼接所有 run 的文本，定位目标子串。
    2. 根据 start/end 位置拆分 runs。
    3. 将目标区间替换为「删除 run」+「插入 run」。
    """
    full_text = para.text
    idx = full_text.find(original_text)
    if idx == -1:
        return False

    runs = para.runs
    if not runs:
        return False

    # 构建 run 边界映射: [(run_index, start_in_full, end_in_full), ...]
    run_boundaries = []
    offset = 0
    for i, run in enumerate(runs):
        run_len = len(run.text)
        run_boundaries.append((i, offset, offset + run_len))
        offset += run_len

    end_idx = idx + len(original_text)

    # 找到受影响的 runs
    # start_run: original_text 起始所在的 run
    # end_run: original_text 结束所在的 run
    start_run = None
    end_run = None
    for i, (ri, rs, re) in enumerate(run_boundaries):
        if rs <= idx < re:
            start_run = i
        if rs < end_idx <= re:
            end_run = i
            break

    if start_run is None or end_run is None:
        return False

    # ---- 拆分并重建 runs ----
    new_runs = []

    for i, run in enumerate(runs):
        ri, rs, re = run_boundaries[i]

        if i < start_run:
            # 完全在目标之前的 run，保持不变
            new_runs.append(run)
            continue

        if i > end_run:
            # 完全在目标之后的 run，保持不变
            new_runs.append(run)
            continue

        # ---- 处理受影响的 run ----
        run_text = run.text

        # 计算 original_text 在此 run 中的局部区间
        local_start = max(0, idx - rs)
        local_end = min(len(run_text), end_idx - rs)

        before = run_text[:local_start]
        deleted = run_text[local_start:local_end]
        after = run_text[local_end:]

        # before 部分 — 保持原样式
        if before:
            new_runs.append(_make_run(para, before, run))

        # 如果是第一个受影响的 run 且有 before 文本，则在此处插入
        # deleted run（红色 + 删除线）— 仅在 start_run 上输出
        if i == start_run and deleted:
            new_runs.append(_make_deleted_run(para, deleted))

        # inserted run（蓝色 + 下划线）— 仅在 start_run 上输出
        if i == start_run:
            new_runs.append(_make_inserted_run(para, suggested_text))

        # 非首尾 run 中间部分被完全消费，不再保留

        # after 部分 — 保持原样式
        if i == end_run and after:
            new_runs.append(_make_run(para, after, run))

    # 替换段落中的 runs
    # 先移除所有旧的 run XML 元素
    p_elem = para._element
    for run in runs:
        r_elem = run._element
        p_elem.remove(r_elem)

    # 追加新的 run XML 元素
    for run in new_runs:
        p_elem.append(run._element)

    return True


def _make_run(para, text: str, source_run):
    """创建一个保持 source_run 样式的 run。"""
    from docx.oxml import OxmlElement

    r_elem = deepcopy(source_run._element)
    # 清除文本
    for t in r_elem.findall(qn("w:t")):
        r_elem.remove(t)

    # 设置新文本
    t_elem = OxmlElement("w:t")
    t_elem.text = text
    t_elem.set(qn("xml:space"), "preserve")
    r_elem.append(t_elem)

    # 重建 Run 对象
    from docx.text.run import Run
    new_run = Run(r_elem, para)
    return new_run


def _make_deleted_run(para, text: str):
    """创建「删除」风格的 run：红色 + 删除线。"""
    from docx.oxml import OxmlElement
    from docx.text.run import Run

    r_elem = OxmlElement("w:r")

    # run properties
    rPr = OxmlElement("w:rPr")

    # 红色
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "FF0000")
    rPr.append(color)

    # 删除线
    strike = OxmlElement("w:strike")
    strike.set(qn("w:val"), "true")
    rPr.append(strike)

    r_elem.append(rPr)

    # 文本
    t_elem = OxmlElement("w:t")
    t_elem.text = text
    t_elem.set(qn("xml:space"), "preserve")
    r_elem.append(t_elem)

    return Run(r_elem, para)


def _make_inserted_run(para, text: str):
    """创建「插入」风格的 run：蓝色 + 下划线。"""
    from docx.oxml import OxmlElement
    from docx.text.run import Run

    r_elem = OxmlElement("w:r")

    # run properties
    rPr = OxmlElement("w:rPr")

    # 蓝色
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    # 下划线
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    r_elem.append(rPr)

    # 文本
    t_elem = OxmlElement("w:t")
    t_elem.text = text
    t_elem.set(qn("xml:space"), "preserve")
    r_elem.append(t_elem)

    return Run(r_elem, para)
